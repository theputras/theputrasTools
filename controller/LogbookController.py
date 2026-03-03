import os
import io
from werkzeug.utils import secure_filename
from PIL import Image
from docx import Document
import time
import re
import html
from docx.shared import Inches, Pt, RGBColor
from flask import send_file, current_app
from connection import get_connection
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE # Tambahkan ini di bagian atas file import lu
import bleach
from PIL import Image
from datetime import datetime
import uuid

# UPLOAD_FOLDER = os.path.join('static', 'uploads', 'logbook')
# os.makedirs(UPLOAD_FOLDER, exist_ok=True)
ALLOWED_TAGS = [
    'p', 'ul', 'ol', 'li', 'strong', 'em', 'u', 'h1', 'h2', 'h3', 'br'
]
ALLOWED_IMAGE_EXT = {'.jpg', '.jpeg', '.png', '.gif', '.webp', '.svg', '.ico', '.avif', '.heic', '.heif'}
ALLOWED_TTD_EXT = {'.png', '.jpg', '.jpeg'}

# --- HELPER: SAVE/DELETE FILE TANDA TANGAN ---
def save_signature_file(file, nim):
    """Simpan file tanda tangan mentor ke disk dengan auto-compress. Return relative path atau None jika gagal."""
    if not file or file.filename == '':
        return None
    
    filename = secure_filename(file.filename)
    _, ext = os.path.splitext(filename)
    ext = ext.lower()
    
    # Validasi ekstensi (hanya PNG)
    if ext not in ALLOWED_TTD_EXT:
        return None
    
    # Validasi MIME type (terima PNG dan JPEG, nanti dikonversi ke PNG)
    if file.content_type not in ('image/png', 'image/jpeg', 'image/jpg'):
        return None
    
    # Folder: static/uploads/logbook/{nim}/ttd/
    ttd_folder = os.path.join(current_app.root_path, 'static', 'uploads', 'logbook', str(nim), 'ttd')
    os.makedirs(ttd_folder, exist_ok=True)
    
    # Nama file: signature-{nim}.png (overwrite jika sudah ada)
    save_path = os.path.join(ttd_folder, f'signature-{nim}.png')
    
    # Auto-compress: resize + optimize PNG otomatis
    try:
        img = Image.open(file)
        # Pertahankan transparency (RGBA)
        if img.mode not in ('RGBA', 'LA'):
            img = img.convert('RGBA')
        
        # Resize jika terlalu besar (max width 500px, proporsional)
        max_w = 500
        if img.width > max_w:
            ratio = max_w / img.width
            new_h = int(img.height * ratio)
            img = img.resize((max_w, new_h), Image.LANCZOS)
        
        # Simpan dengan optimize
        img.save(save_path, format='PNG', optimize=True)
        img.close()
    except Exception as e:
        print(f"Gagal compress TTD, simpan apa adanya: {e}")
        file.seek(0)
        file.save(save_path)
    
    return f"{nim}/ttd/signature-{nim}.png"

def delete_signature_file(ttd_path):
    """Hapus file tanda tangan dari disk."""
    if not ttd_path:
        return
    full_path = os.path.join(current_app.root_path, 'static', 'uploads', 'logbook', ttd_path)
    if os.path.exists(full_path):
        try:
            os.remove(full_path)
        except OSError:
            pass

def compress_and_save_image(image_file, nim, entry_id):
    if not image_file or image_file.filename == '':
        return None
        
    original_filename = secure_filename(image_file.filename)
    
    # Kita cuma butuh ekstensinya aja (misal: ".jpg" atau ".png")
    _, ext = os.path.splitext(original_filename)
    
    # Format baru sesuai request lu: {nim}Image_{entry_id}_{timestamp}{ext}
    # Contoh hasil: 23410100003Image_45_1708150000.jpg
    timestamp = int(time.time() * 1000)
    filename = f"{nim}img_{entry_id}_{timestamp}{ext}"
    
    # Bikin struktur folder dinamis
    nim_folder = os.path.join('static', 'uploads', 'logbook', str(nim), 'imgs')
    os.makedirs(nim_folder, exist_ok=True)
    
    filepath = os.path.join(nim_folder, filename)
    
    img = Image.open(image_file)
    if img.mode in ("RGBA", "P"): 
        img = img.convert("RGB")
    img.thumbnail((800, 800))
    img.save(filepath, optimize=True, quality=60)
    
    # Return path relatif untuk disimpan ke DB
    return f"{nim}/imgs/{filename}"
def clean_html_for_word(html_text):
    """Fungsi sakti merubah HTML Tiptap jadi teks rapi buat Word"""
    if not html_text: return ""
    
    # Ubah list item <li> jadi format bullet
    text = re.sub(r'<li>', r'- ', html_text)
    # Ganti tag akhir paragraf/list jadi baris baru (enter)
    text = re.sub(r'</p>|<br>|</li>|</ul>|</ol>', r'\n', text)
    # Hapus semua tag HTML yang tersisa
    text = re.sub(r'<[^>]+>', '', text)
    # Decode simbol (kayak &amp; jadi &)
    text = html.unescape(text)
    
    # Hapus enter berlebih
    return re.sub(r'\n{3,}', '\n\n', text).strip()
# --- CRUD SETUP LOGBOOK ---
def get_logbooks_by_user(user_id):
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    # Filter berdasarkan user_id
    cursor.execute("SELECT * FROM logbooks WHERE user_id = %s ORDER BY id DESC", (user_id,))
    logbooks = cursor.fetchall()
    conn.close()
    return logbooks

def get_logbook_by_id_and_user(id, user_id):
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT * FROM logbooks WHERE id = %s AND user_id = %s", (id, user_id))
    logbook = cursor.fetchone()
    conn.close()
    return logbook

def get_logbook_id_by_uuid(uuid_str):
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT id FROM logbooks WHERE uuid = %s", (uuid_str,))
    result = cursor.fetchone()
    conn.close()
    return result[0] if result else None

def get_entry_id_by_uuid(uuid_str):
    conn = get_connection()
    cursor = conn.cursor()
    cursor.execute("SELECT id FROM logbook_entries WHERE uuid = %s", (uuid_str,))
    result = cursor.fetchone()
    conn.close()
    return result[0] if result else None

def create_logbook(user_id, data, ttd_path=None):
    conn = get_connection()
    cursor = conn.cursor()
    
    new_uuid = str(uuid.uuid4())[:8]
    
    query = """
    INSERT INTO logbooks (uuid, user_id, fakultas, prodi, nama, nim, nama_mitra, waktu_mulai, waktu_selesai, posisi_magang, nama_mentor, wa_mentor, email_mentor, google_doc_id, google_doc_name, ttd_mentor_path)
    VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s, %s)
    """
    val = (
        new_uuid, user_id, data['fakultas'], data['prodi'], data['nama'], data['nim'], 
        data['nama_mitra'], data['waktu_mulai'], data['waktu_selesai'], 
        data['posisi_magang'], data['nama_mentor'], data['wa_mentor'], data['email_mentor'],
        data.get('google_doc_id'), data.get('google_doc_name'),
        ttd_path
    )
    cursor.execute(query, val)
    
    conn.commit()
    conn.close()
    return new_uuid

def update_logbook(id, data, user_id, ttd_path=None, remove_ttd=False):
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    
    # Jika ada file TTD baru atau remove, hapus file lama dulu
    if ttd_path or remove_ttd:
        cursor.execute("SELECT ttd_mentor_path FROM logbooks WHERE id = %s AND user_id = %s", (id, user_id))
        old = cursor.fetchone()
        if old and old.get('ttd_mentor_path'):
            delete_signature_file(old['ttd_mentor_path'])
        
        # Reset signature approvals
        cursor.execute("DELETE FROM logbook_signatures WHERE logbook_id = %s", (id,))
    
    # Build query dynamically
    fields = """fakultas=%s, prodi=%s, nama=%s, nim=%s, nama_mitra=%s, 
    waktu_mulai=%s, waktu_selesai=%s, posisi_magang=%s, 
    nama_mentor=%s, wa_mentor=%s, email_mentor=%s,
    google_doc_id=%s, google_doc_name=%s"""
    
    val = [
        data['fakultas'], data['prodi'], data['nama'], data['nim'], 
        data['nama_mitra'], data['waktu_mulai'], data['waktu_selesai'], 
        data['posisi_magang'], data['nama_mentor'], data['wa_mentor'], data['email_mentor'],
        data.get('google_doc_id'), data.get('google_doc_name'),
    ]
    
    if ttd_path:
        fields += ", ttd_mentor_path=%s"
        val.append(ttd_path)
    elif remove_ttd:
        fields += ", ttd_mentor_path=NULL"
    
    val.extend([id, user_id])
    
    query = f"UPDATE logbooks SET {fields} WHERE id = %s AND user_id = %s"
    cursor.execute(query, tuple(val))
    conn.commit()
    conn.close()

def delete_logbook(id, user_id):
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    
    # 1. Pastikan logbooknya ada dan milik user tersebut, sekaligus ambil nim-nya
    cursor.execute("SELECT nim FROM logbooks WHERE id = %s AND user_id = %s", (id, user_id))
    logbook = cursor.fetchone()
    
    if logbook:
        nim = logbook['nim']
        
        # 2. Ambil semua path gambar dari tabel logbook_images
        # Join dengan logbook_entries untuk filter by logbook_id
        query_get_images = """
            SELECT i.path 
            FROM logbook_images i
            JOIN logbook_entries e ON i.entry_id = e.id
            WHERE e.logbook_id = %s
        """
        cursor.execute(query_get_images, (id,))
        images = cursor.fetchall()
        
        # 3. Hapus file gambarnya satu per satu secara fisik
        for img in images:
            if img['path']:
                file_path = os.path.join(current_app.root_path, 'static', 'uploads', 'logbook', img['path'])
                if os.path.exists(file_path):
                    try:
                        os.remove(file_path)
                    except Exception as e:
                        print(f"Gagal hapus file {file_path}: {e}")
                    
        # 4. Hapus logbook dari database 
        # (Cascade delete akan otomatis hapus entries & images di DB)
        cursor.execute("DELETE FROM logbooks WHERE id = %s AND user_id = %s", (id, user_id))
        conn.commit()
        
        # 5. BERSIH-BERSIH FOLDER (Hanya dihapus JIKA KOSONG)
        try:
            img_dir = os.path.join('static', 'uploads', 'logbook', str(nim), 'imgs')
            base_dir = os.path.join('static', 'uploads', 'logbook', str(nim))
            
            # os.rmdir hanya akan menghapus folder jika isinya benar-benar kosong
            if os.path.exists(img_dir) and not os.listdir(img_dir):
                os.rmdir(img_dir)
            if os.path.exists(base_dir) and not os.listdir(base_dir):
                os.rmdir(base_dir)
        except OSError:
            # Abaikan jika error (berarti foldernya masih ada isinya dari logbook lain)
            pass

    conn.close()

# --- CRUD ENTRIES (KEGIATAN HARIAN) ---
def get_entries_by_logbook(logbook_id):
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    
    # Ambil entri
    cursor.execute("SELECT * FROM logbook_entries WHERE logbook_id = %s ORDER BY tanggal ASC", (logbook_id,))
    entries = cursor.fetchall()
    
    for entry in entries:
        if entry['tanggal']:
            entry['tanggal_display'] = entry['tanggal'].strftime('%d-%m-%Y')
        else:
            entry['tanggal_display'] = '-'
            
        # UPDATE QUERY: Ambil metadata lengkap
        cursor.execute("""
            SELECT id, path, nama_asli, deskripsi, 
                   tipe_berkas, ukuran_berkas, dimensi, created_at 
            FROM logbook_images 
            WHERE entry_id = %s
        """, (entry['id'],))
        
        images = cursor.fetchall()
        
        # Format ukuran berkas biar enak dibaca (optional logic for backend rendering)
        for img in images:
            if img['ukuran_berkas']:
                # Konversi bytes ke KB/MB simpel
                size_bytes = img['ukuran_berkas']
                if size_bytes < 1024 * 1024:
                    img['ukuran_display'] = f"{round(size_bytes / 1024, 2)} KB"
                else:
                    img['ukuran_display'] = f"{round(size_bytes / (1024 * 1024), 2)} MB"
            else:
                img['ukuran_display'] = "0 KB"

        entry['images'] = images
        
    conn.close()
    return entries

def add_entry(logbook_id, tanggal, aktivitas, deskripsi, files):
    conn = get_connection()
    cursor = conn.cursor()
    
    try:
        # Ambil NIM dari logbook owner untuk struktur folder
        cursor.execute("SELECT nim FROM logbooks WHERE id = %s", (logbook_id,))
        row = cursor.fetchone()
        if not row:
            return False
        nim_user = str(row[0]) # Menggunakan NIM asli

        # 1. Insert Entry
        new_uuid = str(uuid.uuid4())[:8]
        cursor.execute(
            "INSERT INTO logbook_entries (uuid, logbook_id, tanggal, aktivitas, deskripsi) VALUES (%s, %s, %s, %s, %s)",
            (new_uuid, logbook_id, tanggal, aktivitas, deskripsi)
        )
        entry_id = cursor.lastrowid
        
        # 2. Process Files
        for file in files:
            process_and_save_image(file, entry_id, nim_user, cursor)
            
        conn.commit()
        return True
    except Exception as e:
        print(f"Error add_entry: {e}")
        conn.rollback()
        return False
    finally:
        conn.close()
        
def delete_entry(entry_id, logbook_id):
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    
    # Cek apakah entry valid dan milik logbook yang sesuai
    cursor.execute("SELECT id FROM logbook_entries WHERE id = %s AND logbook_id = %s", (entry_id, logbook_id))
    if not cursor.fetchone():
        conn.close()
        return False # Tolak jika hacker mencoba manipulasi URL

    # Ambil semua path gambar dari tabel logbook_images
    cursor.execute("SELECT path FROM logbook_images WHERE entry_id = %s", (entry_id,))
    images = cursor.fetchall()
    
    # Hapus file fisik
    for img in images:
        if img['path']:
            file_path = os.path.join(current_app.root_path, 'static', 'uploads', 'logbook', img['path'])
            if os.path.exists(file_path):
                try:
                    os.remove(file_path)
                except Exception as e:
                    print(f"Gagal hapus file {file_path}: {e}")
            
    # Hapus data dari DB (Cascade LogbookImages)
    cursor.execute("DELETE FROM logbook_entries WHERE id = %s", (entry_id,))
    conn.commit()
    conn.close()
    
def get_entry_by_id(entry_id):
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    
    cursor.execute("SELECT * FROM logbook_entries WHERE id = %s", (entry_id,))
    entry = cursor.fetchone()
    
    if entry:
        if entry['tanggal']:
            entry['tanggal_display'] = entry['tanggal'].strftime('%d-%m-%Y')
            
        # UPDATE QUERY: Ambil metadata lengkap
        cursor.execute("""
            SELECT id, path, nama_asli, deskripsi, 
                   tipe_berkas, ukuran_berkas, dimensi, created_at 
            FROM logbook_images 
            WHERE entry_id = %s
        """, (entry_id,))
        images = cursor.fetchall()
        
        # Format ukuran berkas biar enak dibaca
        for img in images:
            if img['ukuran_berkas']:
                size_bytes = img['ukuran_berkas']
                if size_bytes < 1024 * 1024:
                    img['ukuran_display'] = f"{round(size_bytes / 1024, 2)} KB"
                else:
                    img['ukuran_display'] = f"{round(size_bytes / (1024 * 1024), 2)} MB"
            else:
                img['ukuran_display'] = "0 KB"
        
        entry['images'] = images
        
    conn.close()
    return entry

# Tambahkan parameter logbook_id di sini
def update_entry(entry_id, logbook_id, form_data, files):
    aktivitas = form_data.get('aktivitas')
    deskripsi_raw = form_data.get('deskripsi')

    # SANITASI: Bersihkan HTML dari karakter berbahaya sebelum masuk DB
    deskripsi_clean = bleach.clean(deskripsi_raw, tags=ALLOWED_TAGS, strip=True)
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)

    # ========================================================
    # FIX CRITICAL (IDOR): Validasi kepemilikan entry_id
    # ========================================================
    cursor.execute("""
        SELECT l.nim, e.id as entry_id 
        FROM logbook_entries e 
        JOIN logbooks l ON e.logbook_id = l.id 
        WHERE e.id = %s AND e.logbook_id = %s 
    """, (entry_id, logbook_id))
    
    row = cursor.fetchone()
    
    # Kalau data ga ketemu (berarti hacker nyoba masukin entry_id orang lain), langsung tolak!
    if not row:
        conn.close()
        return False 
        
    nim = row['nim']

    # 1. Update data teks terlebih dahulu
    cursor.execute(
        "UPDATE logbook_entries SET tanggal=%s, aktivitas=%s, deskripsi=%s WHERE id=%s AND logbook_id=%s",
        (form_data.get('tanggal'), aktivitas, deskripsi_clean, entry_id, logbook_id)
    )

    # --- LOGIKA MULTIPLE GAMBAR YANG BENAR ---
    # Ambil array path gambar yang TIDAK DIHAPUS oleh user di form HTML
    retained_paths = form_data.getlist('existing_images') 
    new_images = files.getlist('gambar')

    # Ambil semua gambar lama dari database
    cursor.execute("SELECT path FROM logbook_images WHERE entry_id = %s", (entry_id,))
    old_images = cursor.fetchall()

    # A. Bandingkan dan hapus gambar yang disilang (X) oleh user
    for old_img in old_images:
        if old_img['path'] not in retained_paths:
            # Hapus file fisik dari HDD
            old_path = os.path.join('static', 'uploads', 'logbook', old_img['path'])
            if os.path.exists(old_path):
                try:
                    os.remove(old_path)
                except OSError:
                    pass
            # Hapus dari tabel
            cursor.execute("DELETE FROM logbook_images WHERE entry_id = %s AND path = %s", (entry_id, old_img['path']))

    # B. Tambahkan gambar baru (jika ada)
    for img in new_images:
            if img and img.filename != '':
                new_path = compress_and_save_image(img, nim, entry_id)
                # UPDATE: Simpan nama asli file dan deskripsi default kosong
                cursor.execute(
                    "INSERT INTO logbook_images (entry_id, path, nama_asli, deskripsi) VALUES (%s, %s, %s, %s)",
                    (entry_id, new_path, img.filename, "")
                )

    conn.commit()
    conn.close()
    return True

# --- HELPER FUNCTION UNTUK SAVE IMAGE & METADATA ---


def process_and_save_image(file, entry_id, nim_user, cursor):
    if not file:
        return

    filename = secure_filename(file.filename)
    if filename == '':
        return

    # 1. Generate Nama File Sesuai Request
    # Format: {nim}img_{entry_id}_{timestamp}{ext}
    ext = os.path.splitext(filename)[1].lower() # .jpg, .png
    
    # Validasi ekstensi file
    if ext not in ALLOWED_IMAGE_EXT:
        print(f"File ditolak: ekstensi '{ext}' tidak diizinkan.")
        return
    
    timestamp = int(datetime.now().timestamp() * 1000) # timestamp ms
    new_filename = f"{nim_user}img_{entry_id}_{timestamp}{ext}"
    
    # Path Folder: static/uploads/logbook/{nim}/imgs/
    upload_folder = os.path.join(current_app.root_path, 'static', 'uploads', 'logbook', str(nim_user), 'imgs')
    os.makedirs(upload_folder, exist_ok=True)
    
    file_path = os.path.join(upload_folder, new_filename)
    db_path = f"{nim_user}/imgs/{new_filename}" # Path relative untuk DB

    # 2. Simpan File Fisik
    file.save(file_path)

    # 3. Ambil Metadata Tambahan
    # Ukuran Berkas
    file_size = os.path.getsize(file_path) # Bytes
    
    # Tipe Berkas
    file_type = file.content_type # image/jpeg
    
    # Dimensi (Butuh Pillow)
    dimensi_str = "Unknown"
    try:
        with Image.open(file_path) as img_pil:
            width, height = img_pil.size
            dimensi_str = f"{width}x{height}"
    except Exception as e:
        print(f"Gagal baca dimensi gambar: {e}")

    # 4. Insert ke Database dengan Metadata Lengkap
    cursor.execute("""
        INSERT INTO logbook_images 
        (entry_id, path, nama_asli, deskripsi, tipe_berkas, ukuran_berkas, dimensi) 
        VALUES (%s, %s, %s, %s, %s, %s, %s)
    """, (entry_id, db_path, filename, "", file_type, file_size, dimensi_str))
# --- GENERATE WORD ---
# Tambahin parameter user_id biar pas diconvert dicek dulu ownershipnya
def generate_word(logbook_id, user_id): 
    logbook = get_logbook_by_id_and_user(logbook_id, user_id) 
    entries = get_entries_by_logbook(logbook_id)

    if not logbook:
        return None 

    doc = Document()
    
    # Global font: Times New Roman, 12pt
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    # 1. Judul Utama
    doc.add_heading('Log Book Bulanan Dinamika Industrial Internship', 0)

    # 2. Tabel Identitas (Tanpa Border)
    if logbook:
        # ... (Kode identitas tetap sama seperti sebelumnya) ...
        # [Bagian ini dilewati untuk mempersingkat jawaban]
        start = logbook['waktu_mulai']
        end = logbook['waktu_selesai']
        start_str = start.strftime('%d-%m-%Y') if hasattr(start, 'strftime') else str(start)
        end_str = end.strftime('%d-%m-%Y') if hasattr(end, 'strftime') else str(end)
        identitas = [
            ("Fakultas", ":", logbook.get('fakultas', '-')),
            ("Prodi", ":", logbook.get('prodi', '-')),
            ("Nama", ":", logbook.get('nama', '-')),
            ("Nim", ":", logbook.get('nim', '-')),
            ("Nama Mitra", ":", logbook.get('nama_mitra', '-')),
            ("Waktu Pelaksanaan", ":", f"{start_str} sampai {end_str}"),
            ("Posisi Magang", ":", logbook.get('posisi_magang', '-')),
            ("Nama Mentor", ":", logbook.get('nama_mentor', '-')),
            ("Whatsapp Mentor", ":", logbook.get('wa_mentor', '-')),
            ("Email Mentor", ":", logbook.get('email_mentor', '-'))
        ]
        id_table = doc.add_table(rows=0, cols=3)
        for label, sep, val in identitas:
            row_cells = id_table.add_row().cells
            row_cells[0].text = label
            row_cells[1].text = sep
            row_cells[2].text = str(val)
            row_cells[0].width = Inches(1.5)
            row_cells[1].width = Inches(0.2)
            row_cells[2].width = Inches(4.3)

    # 3. Pengelompokan Kegiatan Berdasarkan Bulan
    months_id = ['', 'Januari', 'Februari', 'Maret', 'April', 'Mei', 'Juni', 'Juli', 'Agustus', 'September', 'Oktober', 'November', 'Desember']
    grouped_entries = {}
    for entry in entries:
        tgl = entry['tanggal']
        month_key = f"{months_id[tgl.month]} {tgl.year}" if tgl else "Belum Diketahui"
        month_only = months_id[tgl.month] if tgl else ""
        if month_key not in grouped_entries:
            grouped_entries[month_key] = {'month_only': month_only, 'data': []}
        grouped_entries[month_key]['data'].append(entry)

    # --- TAMBAHAN: Buka koneksi database untuk ambil gambar ---
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)

    # 4. Render Logbook Per Bulan
    for month_key, group in grouped_entries.items():
        doc.add_paragraph() 
        heading = doc.add_heading('', level=2)
        run_h = heading.add_run(f'Aktivitas Bulan {month_key}')
        run_h.font.color.rgb = RGBColor(255, 0, 0)
        run_h.font.name = 'Times New Roman'
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        table.autofit = False
        table.allow_autofit = False
        
        # Set explicitly column widths
        table.columns[0].width = Inches(0.5)
        table.columns[1].width = Inches(1.5)
        table.columns[2].width = Inches(4.5)
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'No'
        hdr_cells[1].text = 'Aktivitas'
        hdr_cells[2].text = 'Deskripsi Kegiatan'
        
        # Also set the cell widths for the header just in case
        hdr_cells[0].width = Inches(0.5)
        hdr_cells[1].width = Inches(1.5)
        hdr_cells[2].width = Inches(4.5)

        # Isi Tabel
        for idx, entry in enumerate(group['data'], start=1):
            row_cells = table.add_row().cells
            row_cells[0].text = str(idx)
            row_cells[0].width = Inches(0.5)
            
            tgl_str = entry['tanggal'].strftime('%d-%m-%Y') if entry['tanggal'] else '-'
            row_cells[1].text = f"{tgl_str}\n{entry['aktivitas']}"
            row_cells[1].width = Inches(1.5)
            
            row_cells[2].width = Inches(4.5)
            p = row_cells[2].paragraphs[0]
            clean_deskripsi = clean_html_for_word(entry['deskripsi'])
            p.add_run(clean_deskripsi + "\n")
            
            # --- MODIFIKASI: AMBIL MULTIPLE GAMBAR DARI DATABASE ---
            cursor.execute("SELECT path FROM logbook_images WHERE entry_id = %s", (entry['id'],))
            db_images = cursor.fetchall()

# --- MODIFIKASI: AMBIL MULTIPLE GAMBAR + METADATA ---
            # Pastikan select deskripsi dan nama_asli juga
            cursor.execute("SELECT path, nama_asli, deskripsi FROM logbook_images WHERE entry_id = %s", (entry['id'],))
            db_images = cursor.fetchall()

            if db_images:
                # Bikin paragraf container untuk gambar
                p_images = row_cells[2].add_paragraph()
                
                for img in db_images:
                    img_path = os.path.join('static', 'uploads', 'logbook', img['path'])
                    if os.path.exists(img_path):
                        run = p_images.add_run()
                        run.add_picture(img_path, width=Inches(2.0)) # Kecilin dikit biar muat banyak
                        run.add_text("\n") # Enter setelah gambar
                        
                        # Tampilkan Keterangan Gambar jika ada
                        nama = img.get('nama_asli')
                        desc = img.get('deskripsi')
                        
                        if nama or desc:
                            caption_run = p_images.add_run()
                            caption_run.font.size = 90000 # (Ukuran font kecil, approx 7pt)
                            caption_run.italic = True
                            
                            info_text = []
                            if nama: info_text.append(f"[{nama}]")
                            if desc: info_text.append(desc)
                            
                            caption_run.add_text(" ".join(info_text) + "\n\n")
                        else:
                            run.add_text("\n")
            
        doc.add_paragraph() 
        
        # 5. Tambahkan Format Resume & Tanda Tangan
        month_name = group['month_only']
        if month_name:
            # Resume header sebagai paragraf bold biasa (bukan heading berwarna)
            resume_p = doc.add_paragraph()
            resume_run = resume_p.add_run(f'Resume Kegiatan Bulan {month_name}:')
            resume_run.bold = True
            resume_run.font.name = 'Times New Roman'
            resume_run.font.size = Pt(12)
            
            # Ambil resume dari database
            resume_content = get_resume_content(logbook_id, month_key)
            if resume_content:
                clean_resume = clean_html_for_word(resume_content)
                doc.add_paragraph(clean_resume)
            else:
                doc.add_paragraph("...") 
        
        # Cek apakah bulan ini sudah di-approve TTD
        cursor.execute(
            "SELECT is_approved FROM logbook_signatures WHERE logbook_id = %s AND bulan = %s",
            (logbook_id, month_key)
        )
        sig_row = cursor.fetchone()
        month_approved = sig_row and sig_row.get('is_approved', 0) == 1
        
        sig_p = doc.add_paragraph()
        sig_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_label = sig_p.add_run("Disetujui Oleh\n")
        run_label.bold = True
        run_label.font.name = 'Times New Roman'
        run_label.font.size = Pt(12)
        
        # Insert gambar TTD jika bulan sudah approved DAN file TTD ada
        ttd_path = logbook.get('ttd_mentor_path')
        if month_approved and ttd_path:
            # Kalau sudah approved, langsung gambar
            ttd_full_path = os.path.join('static', 'uploads', 'logbook', ttd_path)
            if os.path.exists(ttd_full_path):
                run_ttd = sig_p.add_run()
                try:
                    with Image.open(ttd_full_path) as ttd_img:
                        w, h = ttd_img.size
                        target_w = Inches(1.5)
                        ratio = h / w
                        target_h = Inches(1.5 * ratio)
                    run_ttd.add_picture(ttd_full_path, width=target_w, height=target_h)
                except Exception as e:
                    print(f"Gagal insert TTD ke Word: {e}")
                    r = sig_p.add_run("Tanda Tangan Mentor\n\n\n")
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(12)
                sig_p.add_run("\n")
            else:
                r = sig_p.add_run("Tanda Tangan Mentor\n\n\n")
                r.font.name = 'Times New Roman'
                r.font.size = Pt(12)
        else:
            r = sig_p.add_run("Tanda Tangan Mentor\n\n\n")
            r.font.name = 'Times New Roman'
            r.font.size = Pt(12)
        
        mentor_run = sig_p.add_run(f"{logbook.get('nama_mentor', 'Nama Mentor')}")
        mentor_run.bold = True
        mentor_run.underline = WD_UNDERLINE.SINGLE
        mentor_run.font.name = 'Times New Roman'
        mentor_run.font.size = Pt(12)

    # Tutup koneksi setelah selesai looping
    cursor.close()
    conn.close()

    # Save to Stream
    file_stream = io.BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    current_time = int(time.time())
    
    return send_file(
        file_stream, 
        as_attachment=True, 
        download_name=f"Logbook_{logbook['nim']}_{current_time}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )

# --- CRUD KHUSUS IMAGE (SINGLE) ---

def get_image_by_id(image_id):
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute("SELECT * FROM logbook_images WHERE id = %s", (image_id,))
    image = cursor.fetchone()
    conn.close()
    return image

def delete_single_image(image_id, user_id):
    """Menghapus satu file gambar spesifik dengan validasi User"""
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    
    # 1. Cek Kepemilikan & Ambil Path
    query = """
        SELECT i.id, i.path 
        FROM logbook_images i
        JOIN logbook_entries e ON i.entry_id = e.id
        JOIN logbooks l ON e.logbook_id = l.id
        WHERE i.id = %s AND l.user_id = %s
    """
    cursor.execute(query, (image_id, user_id))
    image = cursor.fetchone()
    
    if image:
        # Hapus file fisik
        file_path = os.path.join('static', 'uploads', 'logbook', image['path'])
        if os.path.exists(file_path):
            try:
                os.remove(file_path)
            except OSError:
                pass 
        
        # Hapus record dari DB
        cursor.execute("DELETE FROM logbook_images WHERE id = %s", (image_id,))
        conn.commit()
        conn.close()
        return True
        
    conn.close()
    return False

def update_image_metadata(image_id, user_id, nama_baru, deskripsi_baru):
    """Update nama dan deskripsi gambar dengan keamanan User ID"""
    conn = get_connection()
    cursor = conn.cursor(dictionary=True) # Pakai dictionary cursor biar enak
    try:
        # 1. Cek Validasi Kepemilikan (PENTING!)
        # Kita join dari images -> entries -> logbooks untuk cek user_id
        check_query = """
            SELECT i.id 
            FROM logbook_images i
            JOIN logbook_entries e ON i.entry_id = e.id
            JOIN logbooks l ON e.logbook_id = l.id
            WHERE i.id = %s AND l.user_id = %s
        """
        cursor.execute(check_query, (image_id, user_id))
        if not cursor.fetchone():
            return False # User mencoba edit gambar orang lain!

        # 2. Eksekusi Update
        cursor.execute(
            "UPDATE logbook_images SET nama_asli = %s, deskripsi = %s WHERE id = %s",
            (nama_baru, deskripsi_baru, image_id)
        )
        conn.commit()
        return True
    except Exception as e:
        print(f"Error update image: {e}")
        return False
    finally:
        conn.close()

def replace_image_file(image_id, user_id, new_file):
    """Replace file fisik gambar yang sudah ada di DB (untuk crop/rotate/resize)"""
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    try:
        # 1. Validasi kepemilikan + ambil path lama & NIM
        query = """
            SELECT i.id, i.path, i.entry_id, l.nim
            FROM logbook_images i
            JOIN logbook_entries e ON i.entry_id = e.id
            JOIN logbooks l ON e.logbook_id = l.id
            WHERE i.id = %s AND l.user_id = %s
        """
        cursor.execute(query, (image_id, user_id))
        image = cursor.fetchone()

        if not image:
            return False, "Unauthorized"

        old_path = image['path']
        nim = image['nim']
        entry_id = image['entry_id']

        # 2. Hapus file fisik lama
        old_file_path = os.path.join('static', 'uploads', 'logbook', old_path)
        if os.path.exists(old_file_path):
            try:
                os.remove(old_file_path)
            except OSError:
                pass

        # 3. Simpan file baru (reuse folder structure)
        timestamp = int(time.time() * 1000)
        new_filename = f"{nim}img_{entry_id}_{timestamp}.jpg"
        nim_folder = os.path.join('static', 'uploads', 'logbook', str(nim), 'imgs')
        os.makedirs(nim_folder, exist_ok=True)

        filepath = os.path.join(nim_folder, new_filename)

        # Compress & save
        img = Image.open(new_file)
        if img.mode in ("RGBA", "P"):
            img = img.convert("RGB")
        # Pertahankan ukuran asli dari canvas crop (sudah di-resize di frontend)
        img.save(filepath, optimize=True, quality=80)

        # 4. Update path di DB
        new_relative_path = f"{nim}/imgs/{new_filename}"
        cursor.execute(
            "UPDATE logbook_images SET path = %s WHERE id = %s",
            (new_relative_path, image_id)
        )
        conn.commit()
        return True, new_relative_path

    except Exception as e:
        print(f"Error replace image: {e}")
        conn.rollback()
        return False, str(e)
    finally:
        conn.close()

# --- CRUD TANDA TANGAN / SIGNATURES ---

def get_available_months(logbook_id):
    """Ambil daftar bulan unik dari logbook_entries untuk logbook tertentu."""
    months_id = ['', 'Januari', 'Februari', 'Maret', 'April', 'Mei', 'Juni', 
                 'Juli', 'Agustus', 'September', 'Oktober', 'November', 'Desember']
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute(
        "SELECT DISTINCT MONTH(tanggal) as bulan, YEAR(tanggal) as tahun FROM logbook_entries WHERE logbook_id = %s ORDER BY tahun ASC, bulan ASC",
        (logbook_id,)
    )
    rows = cursor.fetchall()
    conn.close()
    
    result = []
    for row in rows:
        if row['bulan'] and row['tahun']:
            month_name = months_id[row['bulan']]
            result.append(f"{month_name} {row['tahun']}")
    return result

def get_signatures_by_logbook(logbook_id):
    """Ambil semua data signature approval untuk logbook tertentu, di-merge dengan bulan dari entries."""
    available_months = get_available_months(logbook_id)
    
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute(
        "SELECT bulan, is_approved, approved_at FROM logbook_signatures WHERE logbook_id = %s",
        (logbook_id,)
    )
    sig_rows = cursor.fetchall()
    conn.close()
    
    # Map bulan -> approval data
    sig_map = {}
    for row in sig_rows:
        sig_map[row['bulan']] = {
            'is_approved': row['is_approved'],
            'approved_at': row['approved_at'].strftime('%d-%m-%Y %H:%M') if row['approved_at'] else None
        }
    
    # Merge: setiap bulan yang ada entry-nya punya status approval
    result = []
    for month in available_months:
        data = sig_map.get(month, {'is_approved': 0, 'approved_at': None})
        result.append({
            'bulan': month,
            'is_approved': data['is_approved'],
            'approved_at': data['approved_at']
        })
    
    return result

def approve_signature(logbook_id, bulan, user_id):
    """Approve tanda tangan untuk bulan tertentu. Validasi ownership."""
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    
    # Validasi ownership
    cursor.execute("SELECT id FROM logbooks WHERE id = %s AND user_id = %s", (logbook_id, user_id))
    if not cursor.fetchone():
        conn.close()
        return False
    
    try:
        cursor.execute("""
            INSERT INTO logbook_signatures (logbook_id, bulan, is_approved, approved_at) 
            VALUES (%s, %s, 1, NOW())
            ON DUPLICATE KEY UPDATE is_approved = 1, approved_at = NOW()
        """, (logbook_id, bulan))
        conn.commit()
        return True
    except Exception as e:
        print(f"Error approve signature: {e}")
        return False
    finally:
        conn.close()

def revoke_signature(logbook_id, bulan, user_id):
    """Revoke/cabut tanda tangan untuk bulan tertentu."""
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    
    # Validasi ownership
    cursor.execute("SELECT id FROM logbooks WHERE id = %s AND user_id = %s", (logbook_id, user_id))
    if not cursor.fetchone():
        conn.close()
        return False
    
    try:
        cursor.execute("""
            INSERT INTO logbook_signatures (logbook_id, bulan, is_approved, approved_at) 
            VALUES (%s, %s, 0, NULL)
            ON DUPLICATE KEY UPDATE is_approved = 0, approved_at = NULL
        """, (logbook_id, bulan))
        conn.commit()
        return True
    except Exception as e:
        print(f"Error revoke signature: {e}")
        return False
    finally:
        conn.close()

# --- CRUD RESUME KEGIATAN BULANAN ---

def get_resumes_by_logbook(logbook_id):
    """Ambil semua resume per bulan untuk logbook tertentu, di-merge dengan bulan dari entries."""
    available_months = get_available_months(logbook_id)
    
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute(
        "SELECT bulan, content, updated_at FROM logbook_resumes WHERE logbook_id = %s",
        (logbook_id,)
    )
    resume_rows = cursor.fetchall()
    conn.close()
    
    # Map bulan -> resume data
    resume_map = {}
    for row in resume_rows:
        resume_map[row['bulan']] = {
            'content': row['content'] or '',
            'updated_at': row['updated_at'].strftime('%d-%m-%Y %H:%M') if row['updated_at'] else None
        }
    
    # Merge: setiap bulan yang ada entry-nya punya status resume
    result = []
    for month in available_months:
        data = resume_map.get(month, {'content': '', 'updated_at': None})
        result.append({
            'bulan': month,
            'content': data['content'],
            'updated_at': data['updated_at'],
            'is_filled': bool(data['content'] and data['content'].strip())
        })
    
    return result

def save_resume(logbook_id, bulan, content, user_id):
    """Simpan/update resume untuk bulan tertentu. Validasi ownership."""
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    
    # Validasi ownership
    cursor.execute("SELECT id FROM logbooks WHERE id = %s AND user_id = %s", (logbook_id, user_id))
    if not cursor.fetchone():
        conn.close()
        return False
    
    try:
        cursor.execute("""
            INSERT INTO logbook_resumes (logbook_id, bulan, content) 
            VALUES (%s, %s, %s)
            ON DUPLICATE KEY UPDATE content = VALUES(content)
        """, (logbook_id, bulan, content))
        conn.commit()
        return True
    except Exception as e:
        print(f"Error save resume: {e}")
        return False
    finally:
        conn.close()

def delete_resume(logbook_id, bulan, user_id):
    """Hapus resume untuk bulan tertentu. Validasi ownership."""
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    
    # Validasi ownership
    cursor.execute("SELECT id FROM logbooks WHERE id = %s AND user_id = %s", (logbook_id, user_id))
    if not cursor.fetchone():
        conn.close()
        return False
    
    try:
        cursor.execute("DELETE FROM logbook_resumes WHERE logbook_id = %s AND bulan = %s", (logbook_id, bulan))
        conn.commit()
        return True
    except Exception as e:
        print(f"Error delete resume: {e}")
        return False
    finally:
        conn.close()

def get_resume_content(logbook_id, bulan):
    """Ambil konten resume untuk bulan tertentu (dipakai di generate_word & Google Docs sync)."""
    conn = get_connection()
    cursor = conn.cursor(dictionary=True)
    cursor.execute(
        "SELECT content FROM logbook_resumes WHERE logbook_id = %s AND bulan = %s",
        (logbook_id, bulan)
    )
    row = cursor.fetchone()
    conn.close()
    return row['content'] if row and row['content'] else None